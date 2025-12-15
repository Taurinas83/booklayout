"""
Processador de Manuscritos
Extrai e analisa estrutura de documentos (txt, pdf, docx)
"""

import os
import re
from docx import Document
import PyPDF2
from typing import Dict, List, Any
import zipfile
import xml.etree.ElementTree as ET


class ManuscriptProcessor:
    """Processa manuscritos em diferentes formatos"""
    
    def __init__(self):
        self.chapter_patterns = [
            r'^(CAPÍTULO|CHAPTER|CAP\.?)\s+(\d+)',
            r'^(PARTE|PART)\s+(\d+)',
            r'^#{1,2}\s+(.+)',  # Markdown headers
        ]
        self.section_patterns = [
            r'^(SEÇÃO|SECTION|SEC\.?)\s+(\d+)',
            r'^#{3}\s+(.+)',  # Markdown h3
        ]
    
    def process(self, filepath: str) -> Dict[str, Any]:
        """
        Processa arquivo e retorna estrutura do manuscrito
        """
        ext = os.path.splitext(filepath)[1].lower()
        
        try:
            if ext == '.txt':
                content = self._read_txt(filepath)
            elif ext == '.docx':
                try:
                    content = self._read_docx(filepath)
                except Exception as e:
                    print(f"Erro ao ler DOCX com biblioteca padrão: {e}. Tentando fallback...")
                    content = self._read_docx_fallback(filepath)
            elif ext == '.pdf':
                content = self._read_pdf(filepath)
            else:
                raise ValueError(f"Formato não suportado: {ext}")
        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo {ext}: {str(e)}")
        
        # Analisar estrutura
        structure = self._analyze_structure(content)

    def _read_docx_fallback(self, filepath: str) -> str:
        """
        Lê arquivo DOCX tratando como ZIP e extraindo XML 
        (Fallback seguro se python-docx falhar)
        """
        try:
            with zipfile.ZipFile(filepath) as z:
                xml_content = z.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                # Namespace do Word
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                text_content = []
                for p in root.findall('.//w:p', ns):
                    texts = p.findall('.//w:t', ns)
                    if texts:
                        para_text = ''.join([t.text for t in texts if t.text])
                        if para_text.strip():
                            text_content.append(para_text)
                            
            return '\n\n'.join(text_content)
        except Exception as e:
            raise ValueError(f"Falha crítica no fallback de leitura: {str(e)}")
        
        return {
            'content': content,
            'structure': structure,
            'word_count': len(content.split()),
            'char_count': len(content),
            'file_path': filepath
        }
    
    def _read_txt(self, filepath: str) -> str:
        """Lê arquivo de texto simples"""
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_docx(self, filepath: str) -> str:
        """Lê arquivo Word (.docx)"""
        doc = Document(filepath)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Adicionar texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text)
        
        return '\n\n'.join(content)
    
    def _read_pdf(self, filepath: str) -> str:
        """Lê arquivo PDF"""
        content = []
        
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
        except Exception as e:
            raise ValueError(f"Erro ao ler PDF: {str(e)}")
        
        return '\n\n'.join(content)
    
    def _analyze_structure(self, content: str) -> Dict[str, Any]:
        """
        Analisa estrutura do manuscrito
        Identifica: capítulos, seções, parágrafos, citações
        """
        lines = content.split('\n')
        structure = {
            'chapters': [],
            'sections': [],
            'paragraphs': [],
            'citations': [],
            'images': []
        }
        
        current_chapter = None
        current_section = None
        current_paragraph = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                # Linha vazia - fim de parágrafo
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    structure['paragraphs'].append({
                        'text': para_text,
                        'chapter': current_chapter,
                        'section': current_section,
                        'line_number': i
                    })
                    current_paragraph = []
                continue
            
            # Verificar se é capítulo
            is_chapter = False
            for pattern in self.chapter_patterns:
                match = re.match(pattern, line_stripped, re.IGNORECASE)
                if match:
                    current_chapter = {
                        'title': line_stripped,
                        'line_number': i,
                        'sections': []
                    }
                    structure['chapters'].append(current_chapter)
                    is_chapter = True
                    break
            
            if is_chapter:
                continue
            
            # Verificar se é seção
            is_section = False
            for pattern in self.section_patterns:
                match = re.match(pattern, line_stripped, re.IGNORECASE)
                if match:
                    current_section = {
                        'title': line_stripped,
                        'line_number': i,
                        'chapter': current_chapter['title'] if current_chapter else None
                    }
                    structure['sections'].append(current_section)
                    if current_chapter:
                        current_chapter['sections'].append(current_section)
                    is_section = True
                    break
            
            if is_section:
                continue
            
            # Verificar citações (entre aspas)
            if '"' in line_stripped or "'" in line_stripped:
                citations = re.findall(r'["\']([^"\']+)["\']', line_stripped)
                for citation in citations:
                    structure['citations'].append({
                        'text': citation,
                        'line_number': i,
                        'chapter': current_chapter['title'] if current_chapter else None
                    })
            
            # Adicionar à parágrafo atual
            current_paragraph.append(line_stripped)
        
        # Adicionar último parágrafo se houver
        if current_paragraph:
            para_text = ' '.join(current_paragraph)
            structure['paragraphs'].append({
                'text': para_text,
                'chapter': current_chapter['title'] if current_chapter else None,
                'section': current_section['title'] if current_section else None
            })
        
        return structure
    
    def extract_images(self, filepath: str) -> List[str]:
        """
        Extrai imagens de documentos (DOCX)
        Retorna lista de caminhos das imagens
        """
        images = []
        
        if filepath.endswith('.docx'):
            try:
                doc = Document(filepath)
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        # Salvar imagem temporária
                        image_path = f"temp_image_{len(images)}.png"
                        with open(image_path, 'wb') as f:
                            f.write(image_bytes)
                        images.append(image_path)
            except Exception as e:
                print(f"Erro ao extrair imagens: {str(e)}")
        
        return images
        